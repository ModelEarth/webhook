from flask import Flask, request, jsonify
from docx import Document
import os
import re
from docx.shared import Pt
from flask_mail import Mail, Message
import datetime
app = Flask(__name__)

app.config.from_pyfile('settings.py')

mail = Mail(app)

PLACEHOLDERS = {
    '#name#': 'name',
    '#handle#': 'handle',
    '#team#': 'team',
    '#focus#': 'focus',
    '#goal#': 'goal',
    '#schoolDegree#': 'schoolDegree',
    '#degreeDate#': 'degreeDate',
    '#optDepartment#': 'optDepartment',
    '#optContact#': 'optContact',
    '#workingHours#': 'workingHours',
    '#location#': 'location',
    '#status#': 'status',
    '#github#': 'github',
    '#email#': 'email',
    '#phone#': 'phone',
    '#startDate#': 'startDate',
    '#endDate#': 'endDate',
    '#website#': 'website',
    '#note#': 'note',
    '#jobTitle#': 'jobTitle',
    '#todaysDate#': 'todaysDate',
}

FIELD_ALIASES = {
    'name': ['Name'],
    'handle': ['Handle', 'Programming'],
    'team': ['Team'],
    'focus': ['Focus'],
    'goal': ['UN Goal', 'Goal'],
    'schoolDegree': ['schoolDegree', 'School and Degree Program', 'School Degree', 'Degree'],
    'degreeDate': ['degreeDate', 'Date Degree Completed', 'Degree Date'],
    'optDepartment': ['optDepartment', 'OPT University Department', 'University Department'],
    'optContact': ['optContact', 'OPT University Department Email and/or Phone', 'University Department Email and/or Phone', 'OPT Contact'],
    'workingHours': ['workingHours', "Hours/week you're contributing", 'Working Hours', 'Hours'],
    'location': ['Your Location', 'Location'],
    'status': ['Status'],
    'github': ['Github', 'GitHub', 'Individual GitHub Account', 'Individual GitHub'],
    'email': ['Email', 'Email Address', 'Individual Email'],
    'phone': ['Phone', 'Individual Phone'],
    'startDate': ['startDate', 'StartDate', 'Start Date', 'Starting date'],
    'endDate': ['endDate', 'EndDate', 'End Date'],
    'website': ['Your Website', 'Website'],
    'note': ['Note'],
    'jobTitle': ['Job Title'],
}

LEGACY_INDEXES = {
    'name': ['0'],
    'handle': ['1'],
    'focus': ['2'],
    'goal': ['3'],
    'schoolDegree': ['4'],
    'degreeDate': ['5'],
    'optDepartment': ['6'],
    'optContact': ['7'],
    'workingHours': ['8', '9'],
    'location': ['9', '10'],
    'status': ['10', '11'],
    'github': ['11', '12'],
    'phone': ['12', '14'],
    'email': ['13'],
    'startDate': ['14', '15'],
    'endDate': ['15', '16'],
    'website': ['16', '17'],
    'note': ['17', '18'],
    'jobTitle': ['19'],
}

@app.route("/")
def main():
    return "You've arrived at our Model.Earth webhook test page."

@app.route('/signup', methods=['POST'])
def webhook():
    response = request.json
    doc = Document('files/YourName-ModelEarth-WelcomeLetter.docx')
    answers = build_field_answers(response)
    first_name, full_name = to_camel_case(answers.get('name', ''))
    todays_date = datetime.datetime.now().strftime("%B %-d, %Y")

    github_path = build_github_path(answers.get('github', ''))
    if not answers.get('phone'):
        remove_paragraphs_containing(doc, 'Individual Phone:')
    
    # replace placeholders in paragraphs
    def replace_in_paragraph(paragraph):
        has_changes = False
        
        for placeholder, field in PLACEHOLDERS.items():
            if placeholder in paragraph.text:
                to_replace = ''
                if placeholder == '#todaysDate#':
                    to_replace = todays_date
                elif placeholder == '#name#':
                    to_replace = ' '.join(full_name)
                else:
                    to_replace = answers.get(field, '')
                
                # Replace placeholder while preserving formatting
                for run in paragraph.runs:
                    if placeholder in run.text:
                        # Store original formatting
                        is_bold = run.bold
                        font_name = run.font.name
                        font_size = run.font.size
                        
                        # Replace text
                        run.text = run.text.replace(placeholder, to_replace)
                        
                        # Restore formatting
                        run.bold = is_bold
                        run.font.name = font_name or 'Calibri'
                        run.font.size = font_size or Pt(11)
                        has_changes = True
    
    # Replace in document paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
  
    file_name = f"{''.join(full_name)}-ModelEarth-WelcomeLetter.docx"
    doc_path = os.path.join('/tmp', file_name)
    doc.save(doc_path)

    try:
        send_email(file_name, doc_path, first_name, answers.get('email', ''), github_path)
    except Exception as e:
        return jsonify({'Error': str(e)}), 500
    else:
        return jsonify({'status': 'success', 'documentPath': doc_path}), 200

def build_field_answers(response: dict) -> dict[str, str]:
    answers = {}
    fields_seen = set()
    question_answers = {}
    raw_answers = response.get('namedValues', response) if isinstance(response, dict) else {}

    for key, value in raw_answers.items():
        question = key
        answer = value
        if isinstance(value, dict):
            question = value.get('question', key)
            answer = value.get('answer', '')
        question_answers[normalize_question(question)] = normalize_answer(answer)

    for field, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            normalized_alias = normalize_question(alias)
            if normalized_alias in question_answers:
                answers[field] = question_answers[normalized_alias]
                fields_seen.add(field)
                break

    for field, indexes in LEGACY_INDEXES.items():
        if field in fields_seen or answers.get(field):
            continue
        for index in indexes:
            answer = normalize_answer(raw_answers.get(index, ''))
            if answer:
                answers[field] = answer
                break

    return answers

def normalize_question(question: str) -> str:
    return re.sub(r'[^a-z0-9]+', '', str(question).lower())

def normalize_answer(answer) -> str:
    if isinstance(answer, dict):
        answer = answer.get('answer', '')
    if isinstance(answer, list):
        return ', '.join(str(item).strip() for item in answer if str(item).strip())
    return str(answer).strip() if answer is not None else ''

def build_github_path(github_value: str) -> str:
    if not github_value:
        return ''

    github_value = github_value.strip()
    github_value = re.sub(r'^https?://github\.com/', '', github_value)
    github_value = re.sub(r'^github\.com/', '', github_value)
    github_account = github_value.lstrip('/').split('/')[0]
    return f"https://github.com/{github_account}" if github_account else ''

def remove_paragraphs_containing(doc: Document, text: str) -> None:
    for paragraph in list(doc.paragraphs):
        if text in paragraph.text:
            element = paragraph._element
            element.getparent().remove(element)
            paragraph._p = paragraph._element = None

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    if text in paragraph.text:
                        element = paragraph._element
                        element.getparent().remove(element)
                        paragraph._p = paragraph._element = None

def send_email(file_name: str, doc_path: str, first_name: str, email: str, github_path: str) -> None:
    if not email:
        raise ValueError('Email field is required to send the welcome letter.')

    with app.open_resource('outbound/new-member.txt') as tmp_body:
        msg_body = tmp_body.read().decode('utf-8')
        msg_body = msg_body.replace('[FirstName]', first_name)
        msg_body = msg_body.replace('[githubPath]', github_path)

        msg = Message(
            subject="Welcome to our model.earth team!",
            recipients=[email],
            cc=[app.config['MAIL_DEFAULT_SENDER']]
        )
        msg.body = msg_body
        with app.open_resource(doc_path) as f:
            msg.attach(
                file_name, 
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                f.read()
            )
            mail.send(msg)

def to_camel_case(name: str) -> list[str, list[str]]:
    parts = name.split()
    if not parts:
        return '', []
    camel_case_parts = [part.capitalize() for part in parts]
    return camel_case_parts[0], camel_case_parts

if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=8080)
